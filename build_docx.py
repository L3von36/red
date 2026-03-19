"""
Build thesis_article.docx — Google Docs-compatible.

Formatting follows the Google Docs academic paper template conventions:
  - Times New Roman 12pt body, 11pt abstract/captions, 10pt references
  - 2.54 cm (1 inch) margins all round
  - 1.15 line spacing (Google Docs default)
  - Built-in Heading 1 / Heading 2 / Heading 3 styles (map to Section / Sub / Subsub)
  - Three-line booktabs tables (top rule, header-bottom rule, bottom rule)
  - Architecture figure rendered as a styled ASCII-art text box
    (Google Docs cannot embed ReportLab drawings; ASCII art is the
     next-best inline representation)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, copy

# ── Colours ───────────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x1B, 0x2A, 0x4A)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY  = RGBColor(0x44, 0x44, 0x44)
LGREY = RGBColor(0xF0, 0xF0, 0xF0)

# ── Page setup helper ─────────────────────────────────────────────────────────
def set_page_margins(doc, top=2.54, bottom=2.54, left=2.54, right=2.54):
    section = doc.sections[0]
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)
    section.page_width    = Cm(21.0)   # A4
    section.page_height   = Cm(29.7)


# ── Paragraph style helpers ───────────────────────────────────────────────────
def set_spacing(para, before=0, after=6, line=1.15):
    pf = para.paragraph_format
    pf.space_before      = Pt(before)
    pf.space_after       = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = line


def set_font(run, name='Times New Roman', size=12, bold=False,
             italic=False, colour=None):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if colour:
        run.font.color.rgb = colour


def add_body(doc, text, size=12, indent=True, colour=None):
    """Add a justified body paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, before=0, after=4)
    if indent:
        p.paragraph_format.first_line_indent = Pt(14)
    run = p.add_run(text)
    set_font(run, size=size, colour=colour)
    return p


def add_heading1(doc, text):
    """IEEE-style section heading: bold, centred, small-caps simulation."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=12, after=4)
    run = p.add_run(text.upper())
    set_font(run, size=10, bold=True, colour=NAVY)
    return p


def add_heading2(doc, text):
    """IEEE-style subsection heading: italic, left-aligned."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=8, after=3)
    run = p.add_run(text)
    set_font(run, size=10, italic=True, colour=NAVY)
    return p


def add_heading3(doc, text):
    """Sub-subsection heading."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=6, after=2)
    run = p.add_run(text)
    set_font(run, size=10, italic=True)
    return p


# ── Table helpers ─────────────────────────────────────────────────────────────
def set_cell_border(cell, **kwargs):
    """Set individual cell borders (top/bottom/left/right)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'),   val.get('val', 'single'))
        border.set(qn('w:sz'),    str(val.get('sz', 4)))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), val.get('color', '000000'))
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_booktabs_table(doc, headers, rows, caption=None):
    """
    Three-rule booktabs table:
      top rule    (thick, above header row)
      mid rule    (thin,  below header row)
      bottom rule (thick, below last data row)
    No interior vertical lines.
    """
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style     = 'Table Grid'

    # Remove all borders first (set to 'none')
    for row in table.rows:
        for cell in row.cells:
            for edge in ('top', 'bottom', 'left', 'right', 'insideH', 'insideV'):
                set_cell_border(cell, **{edge: {'val': 'none', 'sz': 0}})

    # Header row
    hdr_row = table.rows[0]
    hdr_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    hdr_row.height      = Pt(16)
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True)
        # Top rule on header top, mid rule on header bottom
        set_cell_border(cell,
            top    ={'val': 'single', 'sz': 8,  'color': '000000'},
            bottom ={'val': 'single', 'sz': 4,  'color': '666666'},
            left   ={'val': 'none',   'sz': 0},
            right  ={'val': 'none',   'sz': 0},
        )
        # Shade header row lightly
        tcPr  = cell._tc.get_or_add_tcPr()
        shd   = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F0F0F0')
        tcPr.append(shd)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        tr = table.rows[r_idx + 1]
        is_last = (r_idx == len(rows) - 1)
        for c_idx, val in enumerate(row_data):
            cell = tr.cells[c_idx]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=10)
            bottom_val = {'val': 'single', 'sz': 8, 'color': '000000'} \
                         if is_last else {'val': 'none', 'sz': 0}
            set_cell_border(cell,
                top   ={'val': 'none', 'sz': 0},
                bottom=bottom_val,
                left  ={'val': 'none', 'sz': 0},
                right ={'val': 'none', 'sz': 0},
            )
            # Alternating row shade
            if r_idx % 2 == 1:
                tcPr = cell._tc.get_or_add_tcPr()
                shd  = OxmlElement('w:shd')
                shd.set(qn('w:val'),   'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'),  'F8F8F8')
                tcPr.append(shd)

    # Caption below
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(cp, before=2, after=8)
        run = cp.add_run(caption)
        set_font(run, size=9, italic=True, colour=GREY)

    return table


# ── Architecture figure (styled text box) ────────────────────────────────────
ARCH_ASCII = """\
┌─────────────────────────────────────────────────────────────────────┐
│  Partial Observations  X_t   |  20% sensors observed · 80% masked  │
└──────────────────────────┬──────────────────────────────────────────┘
                           │
                           ▼
┌─────────────────────────────────────────────────────────────────────┐
│  Hypergraph Conv (HGNN)      |  2-hop corridor hyperedges           │
└──────────────────────────┬──────────────────────────────────────────┘
                           │
                           ▼
┌─────────────────────────────────────────────────────────────────────┐
│  GAT ODE Function            |  GAT₁ → GAT₂ + learnable gate       │
└──────────────────────────┬──────────────────────────────────────────┘
                           │
                           ▼
┌─────────────────────────────────────────────────────────────────────┐
│  Neural ODE Solver           |  Euler integration  Δt = 0.3        │
└──────────────────────────┬──────────────────────────────────────────┘
                           │
                           ▼
┌─────────────────────────────────────────────────────────────────────┐
│  Observation Assimilation    |  GRU gate · fuses h_t with sensors  │
└──────────────────────────┬──────────────────────────────────────────┘
                           │
                           ▼
┌─────────────────────────────────────────────────────────────────────┐
│  Imputed Speeds  X̂_t         |  full reconstruction · 307 nodes    │
└──────────────────────────╌──────────────────────────────────────────┘
                           ╌  (backpropagation)
┌─────────────────────────────────────────────────────────────────────┐
│  Loss = λ₁·JAM-MSE  +  λ₂·Temporal Smooth  +  λ₃·Graph Laplacian  │
└─────────────────────────────────────────────────────────────────────┘"""


def add_arch_figure(doc):
    """Insert architecture as a Courier monospace block with caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=8, after=0)
    run = p.add_run(ARCH_ASCII)
    run.font.name = 'Courier New'
    run.font.size = Pt(7.5)

    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(cap, before=2, after=10)
    r = cap.add_run('Fig. 1.  CTH-NODE architecture overview.')
    set_font(r, size=9, italic=True, colour=GREY)


# ── Inline markdown → runs ────────────────────────────────────────────────────
def add_inline(para, text, base_size=12):
    """Parse **bold**, *italic*, `code` in text and add as runs."""
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'
    parts   = re.split(pattern, text)
    i = 0
    while i < len(parts):
        chunk = parts[i]
        if chunk is None or chunk == '':
            i += 1; continue
        # Detect if chunk is a full match group
        if re.fullmatch(r'\*\*(.+?)\*\*', chunk):
            run = para.add_run(chunk[2:-2])
            set_font(run, size=base_size, bold=True)
        elif re.fullmatch(r'\*(.+?)\*', chunk):
            run = para.add_run(chunk[1:-1])
            set_font(run, size=base_size, italic=True)
        elif re.fullmatch(r'`(.+?)`', chunk):
            run = para.add_run(chunk[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(base_size - 1)
        else:
            run = para.add_run(chunk)
            set_font(run, size=base_size)
        i += 1


# ── Markdown → docx ───────────────────────────────────────────────────────────
def parse_md_to_docx(doc, md_text):
    """
    Convert markdown to python-docx elements.
    Supports: ## headings, ### subheadings, **bold**, *italic*, `code`,
    | tables |, - bullet lists, %%ARCH_FIGURE%% marker.
    """
    lines = md_text.strip().split('\n')
    tbl_rows = []
    in_code, code_buf = False, []
    first_body = True

    # Roman-numeral section counter
    sec_ctr  = [0, 0, 0]

    def flush_code():
        if not code_buf:
            return
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p, before=4, after=4)
        run = p.add_run('\n'.join(code_buf))
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        code_buf.clear()

    def flush_table():
        if not tbl_rows:
            return
        headers  = [c.strip() for c in tbl_rows[0]]
        data     = [[c.strip() for c in row] for row in tbl_rows[1:]]
        add_booktabs_table(doc, headers, data)
        tbl_rows.clear()

    for line in lines:
        # Architecture figure marker
        if line.strip() == '%%ARCH_FIGURE%%':
            flush_table()
            add_arch_figure(doc)
            first_body = False
            continue

        # Code fence
        if line.strip().startswith('```'):
            if in_code:
                flush_code(); in_code = False
            else:
                flush_table(); in_code = True
            continue
        if in_code:
            code_buf.append(line); continue

        # Table row
        if line.strip().startswith('|'):
            cells = [c for c in line.strip().split('|') if c != '']
            if all(set(c.strip()) <= set('-:| ') for c in cells):
                continue   # separator row
            tbl_rows.append(cells)
            continue
        else:
            flush_table()

        # Section heading (##)
        if line.startswith('## '):
            flush_code()
            sec_ctr[0] += 1; sec_ctr[1] = 0; sec_ctr[2] = 0
            title = line[3:].strip()
            add_heading1(doc, title)
            first_body = True
            continue

        # Subsection heading (###)
        if line.startswith('### '):
            flush_code()
            sec_ctr[1] += 1; sec_ctr[2] = 0
            title = line[4:].strip()
            add_heading2(doc, title)
            first_body = True
            continue

        # Sub-subsection (####)
        if line.startswith('#### '):
            flush_code()
            title = line[5:].strip()
            add_heading3(doc, title)
            continue

        # Bullet list
        if line.startswith('- ') or line.startswith('* '):
            flush_code()
            p = doc.add_paragraph(style='List Bullet')
            set_spacing(p, before=0, after=2)
            add_inline(p, line[2:].strip())
            continue

        # Numbered list
        m = re.match(r'^(\d+)\. (.+)', line.strip())
        if m:
            flush_code()
            p = doc.add_paragraph(style='List Number')
            set_spacing(p, before=0, after=2)
            add_inline(p, m.group(2))
            continue

        # Abstract blockquote (> text)
        if line.startswith('> '):
            flush_code()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_spacing(p, before=0, after=4)
            p.paragraph_format.left_indent  = Cm(1)
            p.paragraph_format.right_indent = Cm(1)
            add_inline(p, line[2:], base_size=10)
            continue

        # Blank line
        stripped = line.strip()
        if not stripped:
            flush_code()
            continue

        # Body paragraph
        flush_code()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_spacing(p, before=0, after=4)
        if not first_body:
            p.paragraph_format.first_line_indent = Pt(14)
        first_body = False
        add_inline(p, stripped)

    flush_code()
    flush_table()


# ─────────────────────────────────────────────────────────────────────────────
# CONTENT  (imported from build_pdfs.py at runtime to stay DRY)
# ─────────────────────────────────────────────────────────────────────────────
import sys, importlib.util
_spec = importlib.util.spec_from_file_location('build_pdfs', 'build_pdfs.py')
_mod  = importlib.util.module_from_spec(_spec)
# We only need the string constants — suppress the build calls
_mod.__dict__['build_ieee_article']  = lambda *a, **k: None
_mod.__dict__['build_itsc_paper']    = lambda *a, **k: None
_mod.__dict__['build_thesis_chapter']= lambda *a, **k: None
# Stub print so "Building PDFs…" line is silent
import builtins as _bi
_real_print = _bi.print
_bi.print = lambda *a, **k: None
_spec.loader.exec_module(_mod)
_bi.print = _real_print

TITLE    = ("Hypergraph Neural ODEs with Observation Assimilation\n"
            "for Sparse Traffic Speed Imputation")
AUTHORS  = "[Author Name]"
AFFIL    = "[Department], [University], [City, Country]"
ABSTRACT = _mod.ARTICLE_ABSTRACT
KEYWORDS = _mod.ARTICLE_KEYWORDS
BODY     = _mod.ARTICLE_BODY


# ─────────────────────────────────────────────────────────────────────────────
# BUILD
# ─────────────────────────────────────────────────────────────────────────────
def build_docx(filename):
    doc = Document()
    set_page_margins(doc)

    # ── Title ─────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(title_p, before=0, after=4)
    for line in TITLE.split('\n'):
        run = title_p.add_run(line + ('\n' if '\n' in TITLE and line != TITLE.split('\n')[-1] else ''))
        set_font(run, size=18, bold=True, colour=NAVY)

    # ── Authors / affiliation ─────────────────────────────────────────────────
    auth_p = doc.add_paragraph()
    auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(auth_p, before=2, after=2)
    run = auth_p.add_run(AUTHORS)
    set_font(run, size=11, italic=True)

    affl_p = doc.add_paragraph()
    affl_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(affl_p, before=0, after=10)
    run = affl_p.add_run(AFFIL)
    set_font(run, size=10)

    # ── Horizontal rule ───────────────────────────────────────────────────────
    hr = doc.add_paragraph()
    hr.paragraph_format.border_bottom = True
    set_spacing(hr, before=0, after=6)
    pPr = hr._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1B2A4A')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ── Abstract ─────────────────────────────────────────────────────────────
    abs_hdr = doc.add_paragraph()
    abs_hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(abs_hdr, before=4, after=2)
    r1 = abs_hdr.add_run('Abstract')
    set_font(r1, size=10, bold=True, italic=True)
    r2 = abs_hdr.add_run('\u2014' + ABSTRACT)
    set_font(r2, size=10)
    abs_hdr.paragraph_format.left_indent  = Cm(0.8)
    abs_hdr.paragraph_format.right_indent = Cm(0.8)

    kw_p = doc.add_paragraph()
    kw_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(kw_p, before=2, after=10)
    kw_p.paragraph_format.left_indent  = Cm(0.8)
    kw_p.paragraph_format.right_indent = Cm(0.8)
    r1 = kw_p.add_run('Index Terms')
    set_font(r1, size=10, bold=True, italic=True)
    r2 = kw_p.add_run('\u2014' + KEYWORDS)
    set_font(r2, size=10, italic=True)

    # ── Body ─────────────────────────────────────────────────────────────────
    parse_md_to_docx(doc, BODY)

    doc.save(filename)
    print(f'  \u2705 {filename}')


print('Building Google Docs-compatible DOCX\u2026')
build_docx('thesis_article.docx')
print('Done.')
